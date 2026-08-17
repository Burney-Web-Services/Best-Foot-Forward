"""Tests for DOCX formatting helpers and preferences."""
import os
import sqlite3
import tempfile
from pathlib import Path
from unittest.mock import patch

from docx.shared import RGBColor

import pytest

# We'll test the functions by importing them after setting up a test DB
from best_foot_forward import docx_helpers


class TestGetAccentColor:
    """Test loading accent color preference from DB."""

    def test_default_teal_when_db_missing(self):
        """When DB is missing, use default teal."""
        with tempfile.TemporaryDirectory() as tmpdir:
            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                color = docx_helpers.get_accent_color()
                # Default teal: RGB(0x50, 0x93, 0x8A)
                assert str(color) == str(RGBColor(0x50, 0x93, 0x8A))

    def test_reads_from_db_hex_color(self):
        """When DB exists, read the accent_color_hex value."""
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "data" / "best_foot_forward.db"
            db_path.parent.mkdir(parents=True, exist_ok=True)

            # Create a minimal DB with document_prefs
            conn = sqlite3.connect(db_path)
            conn.execute("""
                CREATE TABLE document_prefs (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    accent_color_hex TEXT NOT NULL DEFAULT '50938A',
                    font_name TEXT NOT NULL DEFAULT 'Calibri'
                )
            """)
            conn.execute("INSERT INTO document_prefs (id, accent_color_hex, font_name) VALUES (1, 'FF0000', 'Arial')")
            conn.commit()
            conn.close()

            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                color = docx_helpers.get_accent_color()
                # Red: RGB(0xFF, 0x00, 0x00)
                assert str(color) == str(RGBColor(0xFF, 0x00, 0x00))

    def test_default_when_table_empty(self):
        """When document_prefs table exists but is empty, use default."""
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "data" / "best_foot_forward.db"
            db_path.parent.mkdir(parents=True, exist_ok=True)

            conn = sqlite3.connect(db_path)
            conn.execute("""
                CREATE TABLE document_prefs (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    accent_color_hex TEXT NOT NULL DEFAULT '50938A',
                    font_name TEXT NOT NULL DEFAULT 'Calibri'
                )
            """)
            conn.commit()
            conn.close()

            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                color = docx_helpers.get_accent_color()
                # Should fall back to default teal
                assert str(color) == str(RGBColor(0x50, 0x93, 0x8A))


class TestGetFontName:
    """Test loading font preference from DB."""

    def test_default_calibri_when_db_missing(self):
        """When DB is missing, use default Calibri."""
        with tempfile.TemporaryDirectory() as tmpdir:
            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                font_name = docx_helpers.get_font_name()
                assert font_name == "Calibri"

    def test_reads_from_db_font_name(self):
        """When DB exists, read the font_name value."""
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "data" / "best_foot_forward.db"
            db_path.parent.mkdir(parents=True, exist_ok=True)

            conn = sqlite3.connect(db_path)
            conn.execute("""
                CREATE TABLE document_prefs (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    accent_color_hex TEXT NOT NULL DEFAULT '50938A',
                    font_name TEXT NOT NULL DEFAULT 'Calibri'
                )
            """)
            conn.execute("INSERT INTO document_prefs (id, accent_color_hex, font_name) VALUES (1, '50938A', 'Arial')")
            conn.commit()
            conn.close()

            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                font_name = docx_helpers.get_font_name()
                assert font_name == "Arial"

    def test_default_when_table_empty(self):
        """When document_prefs table exists but is empty, use default."""
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "data" / "best_foot_forward.db"
            db_path.parent.mkdir(parents=True, exist_ok=True)

            conn = sqlite3.connect(db_path)
            conn.execute("""
                CREATE TABLE document_prefs (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    accent_color_hex TEXT NOT NULL DEFAULT '50938A',
                    font_name TEXT NOT NULL DEFAULT 'Calibri'
                )
            """)
            conn.commit()
            conn.close()

            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                font_name = docx_helpers.get_font_name()
                assert font_name == "Calibri"


class TestFontFunction:
    """Test the font() helper with preference support."""

    def test_font_uses_preference_when_not_specified(self):
        """When font_name is not specified, use the preference."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Test")

        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "data" / "best_foot_forward.db"
            db_path.parent.mkdir(parents=True, exist_ok=True)

            conn = sqlite3.connect(db_path)
            conn.execute("""
                CREATE TABLE document_prefs (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    accent_color_hex TEXT NOT NULL DEFAULT '50938A',
                    font_name TEXT NOT NULL DEFAULT 'Calibri'
                )
            """)
            conn.execute("INSERT INTO document_prefs (id, accent_color_hex, font_name) VALUES (1, '50938A', 'Georgia')")
            conn.commit()
            conn.close()

            with patch("best_foot_forward.docx_helpers._ROOT", tmpdir):
                docx_helpers.font(r, size=12, bold=True)
                # Font name should come from DB preference
                assert r.font.name == "Georgia"
                assert r.font.size.pt == 12
                assert r.bold is True

    def test_font_explicit_font_name_overrides_preference(self):
        """When font_name is explicitly specified, use it regardless of DB."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Test")

        docx_helpers.font(r, size=12, font_name="Arial")
        assert r.font.name == "Arial"
