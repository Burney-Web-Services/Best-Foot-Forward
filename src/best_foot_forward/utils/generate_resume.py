"""
Reads resume content from resume_data.py and generates a formatted .docx file.
Contact info and education are read from SQLite. Output path is the directory
containing the source JD file (JD_FILE_PATH in resume_data.py).
Run with: python3 src/best_foot_forward/utils/generate_resume.py
"""

import sys as _sys, os as _os
_sys.path.insert(0, (_os.environ.get('BFF_DATA_DIR') or _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '../../../data'))) + '/session')  # data/session/ for session files
_sys.path.insert(0, _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '..')))        # src/best_foot_forward/
del _sys, _os

import os
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_helpers import font, get_accent_color, add_border

from db import (get_contact, get_education, get_conn, register_file, resolve_jd_path,
                resolve_or_create_jd)
from utils.audit_log import log_event
from resume_data import COMPANY, ROLE, JD_FILE_PATH, RESUME
JD_FILE_PATH = resolve_jd_path(JD_FILE_PATH)
try:
    from resume_data import LOCATION_OVERRIDE
except ImportError:
    LOCATION_OVERRIDE = None
try:
    from resume_data import SOURCE_APPLICATION_ID
except ImportError:
    SOURCE_APPLICATION_ID = None
try:
    from resume_data import TAILORING_NOTES
except ImportError:
    TAILORING_NOTES = None

from track_bullets import record_bullet_selections

# ── Formatting helpers ────────────────────────────────────────────────────────

def sp(para, before=0, after=2):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def section_header(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    font(r, bold=True, color=get_accent_color())
    sp(p, before=6, after=2)
    add_border(p)

def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Pt(-10)
    r = p.add_run(f'•  {text}')
    font(r)
    sp(p, before=0, after=1)

def render_resume_as_text(resume, contact_info, location_override=None, education=None):
    lines = []
    lines.append(contact_info['name'].upper())
    contact_parts = [contact_info.get('phone', ''), contact_info.get('email', ''), location_override or contact_info.get('location', '')]
    contact_line = '  |  '.join(part for part in contact_parts if part)
    lines.append(contact_line)
    lines.append('')

    if resume.get('summary'):
        lines.append('SUMMARY')
        lines.append(resume['summary'])
        lines.append('')

    if resume.get('skills'):
        lines.append('SKILLS')
        for skill in resume['skills']:
            lines.append(f"{skill['label']} {skill['content']}")
        lines.append('')

    if resume.get('experience'):
        lines.append('EXPERIENCE')
        for job in resume['experience']:
            location = f", {job['location']}" if job.get('location') else ''
            dates = f"  |  {job['dates']}" if job.get('dates') else ''
            lines.append(f"{job['employer']}{location}{dates}")
            for role_entry in job['roles']:
                lines.append(f"  {role_entry['title']}")
                for b in role_entry['bullets']:
                    text = b['text'] if isinstance(b, dict) else b
                    lines.append(f"    • {text}")
        lines.append('')

    if education:
        lines.append('EDUCATION')
        for edu in education:
            lines.append(f"{edu['institution']}, {edu['location']}")
            lines.append(f"  {edu['degree']}")
        lines.append('')

    if resume.get('additional_experience'):
        lines.append('ADDITIONAL EXPERIENCE')
        for job in resume['additional_experience']:
            location = f", {job['location']}" if job.get('location') else ''
            dates = f"  |  {job['dates']}" if job.get('dates') else ''
            lines.append(f"{job['employer']}{location}{dates}")
            for role_entry in job['roles']:
                lines.append(f"  {role_entry['title']}")
                for b in role_entry['bullets']:
                    lines.append(f"    • {b}")
        lines.append('')

    return '\n'.join(lines)

# ── Data ──────────────────────────────────────────────────────────────────────

CONTACT_INFO = get_contact()
EDUCATION    = get_education()

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(CONTACT_INFO['name'].upper())
font(r, size=16, bold=True, color=get_accent_color())
sp(p, after=2)

# Contact line
contact_parts = [CONTACT_INFO.get('phone', ''), CONTACT_INFO.get('email', ''), LOCATION_OVERRIDE or CONTACT_INFO.get('location', '')]
contact_line = '  |  '.join(part for part in contact_parts if part)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(contact_line)
font(r)
sp(p, after=6)

# Summary
section_header(doc, 'SUMMARY')
p = doc.add_paragraph()
r = p.add_run(RESUME['summary'])
font(r)
sp(p, after=2)

# Skills
section_header(doc, 'SKILLS')
for skill in RESUME['skills']:
    p = doc.add_paragraph()
    r1 = p.add_run(skill['label'] + ' ')
    font(r1, bold=True)
    r2 = p.add_run(skill['content'])
    font(r2)
    sp(p, before=1, after=1)

# Experience
section_header(doc, 'EXPERIENCE')
for job in RESUME['experience']:
    p = doc.add_paragraph()
    r1 = p.add_run(job['employer'])
    font(r1, bold=True)
    location = f", {job['location']}" if job.get('location') else ''
    r2 = p.add_run(f"{location}  |  {job['dates']}")
    font(r2)
    sp(p, before=5, after=1)

    for i, role_entry in enumerate(job['roles']):
        p = doc.add_paragraph()
        r = p.add_run(role_entry['title'])
        font(r, italic=True)
        sp(p, before=0 if i == 0 else 3, after=1)
        for b in role_entry['bullets']:
            bullet(doc, b['text'] if isinstance(b, dict) else b)

# Education
section_header(doc, 'EDUCATION')
for edu in EDUCATION:
    p = doc.add_paragraph()
    r = p.add_run(f"{edu['institution']}, {edu['location']}")
    font(r, bold=True)
    sp(p, before=2, after=1)
    p = doc.add_paragraph()
    r = p.add_run(edu['degree'])
    font(r)
    sp(p, after=3)

# Additional Experience (optional)
if RESUME.get('additional_experience'):
    section_header(doc, 'ADDITIONAL EXPERIENCE')
    for job in RESUME['additional_experience']:
        p = doc.add_paragraph()
        r1 = p.add_run(job['employer'])
        font(r1, bold=True)
        location = f", {job['location']}" if job.get('location') else ''
        dates = f"  |  {job['dates']}" if job.get('dates') else ''
        r2 = p.add_run(f"{location}{dates}")
        font(r2)
        sp(p, before=5, after=1)

        for i, role_entry in enumerate(job['roles']):
            p = doc.add_paragraph()
            r = p.add_run(role_entry['title'])
            font(r, italic=True)
            sp(p, before=0 if i == 0 else 3, after=1)
            for b in role_entry['bullets']:
                bullet(doc, b)

# ── Save ──────────────────────────────────────────────────────────────────────

output_dir = os.path.dirname(JD_FILE_PATH)
os.makedirs(output_dir, exist_ok=True)
# sanitize path separators in the role/company (e.g. "Data/Integration") so they
# don't turn the filename into a subdirectory
filename = f"{CONTACT_INFO['name']} Resume - {COMPANY} {ROLE}.docx".replace("/", "-").replace("\\", "-")
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved to: {filepath}")

# Generate plain-text version with same filename stem
txt_filename = filename.replace(".docx", ".txt")
txt_filepath = os.path.join(output_dir, txt_filename)
txt_content = render_resume_as_text(RESUME, CONTACT_INFO, LOCATION_OVERRIDE, EDUCATION)
with open(txt_filepath, 'w', encoding='utf-8') as f:
    f.write(txt_content)
print(f"Saved to: {txt_filepath}")

# ── Persist application to SQLite ─────────────────────────────────────────────

conn = get_conn()
# Resolve the JD row, creating it if this job was never evaluated. Tailoring
# straight from a URL/paste (no evaluate-job pass) means nothing has written a
# jds row yet, and an application with jd_id NULL is orphaned from every report
# that joins through jds.
jd_id, _jd_action = resolve_or_create_jd(conn, COMPANY, ROLE, JD_FILE_PATH)
if _jd_action == "created":
    print(f"[generate_resume] No jds row for {COMPANY} - {ROLE}; inserted one (jd_id={jd_id}).")

conn.execute(
    "UPDATE jds SET output_dir = ? WHERE id = ? AND (output_dir IS NULL OR output_dir != ?)",
    (output_dir, jd_id, output_dir)
)

# jd_id is never None now, so this is a plain equality match. It previously used
# `IS ?`, which with a NULL jd_id matched *any* orphaned application and would
# overwrite an unrelated company's resume_summary/tailoring_notes.
app = conn.execute(
    "SELECT id FROM applications WHERE jd_id = ?",
    (jd_id,)
).fetchone()

if not app:
    now = datetime.now(timezone.utc).isoformat()
    cur = conn.execute(
        """INSERT INTO applications
           (jd_id, created_at, applied_at, resume_summary, stage, source_application_id, tailoring_notes)
           VALUES (?, ?, ?, ?, 'application', ?, ?)""",
        (jd_id, now, now, RESUME['summary'], SOURCE_APPLICATION_ID, TAILORING_NOTES)
    )
    app_id = cur.lastrowid
else:
    app_id = app["id"]
    conn.execute(
        """UPDATE applications SET resume_summary = ?,
           stage = COALESCE(stage, 'application'),
           source_application_id = COALESCE(source_application_id, ?),
           tailoring_notes = COALESCE(tailoring_notes, ?)
           WHERE id = ?""",
        (RESUME['summary'], SOURCE_APPLICATION_ID, TAILORING_NOTES, app_id)
    )

record_bullet_selections(conn, app_id, RESUME)
conn.commit()
conn.close()

log_event("resume-tailor", "generate_resume", application_id=app_id, jd_id=jd_id, company=COMPANY, role=ROLE, filepath=filepath)

register_file(
    filepath, "resume",
    f"Resume tailored for {COMPANY} – {ROLE}",
    jd_id=jd_id, application_id=app_id,
)

register_file(
    txt_filepath, "resume",
    f"Resume tailored for {COMPANY} – {ROLE} (plain text)",
    jd_id=jd_id, application_id=app_id,
)
