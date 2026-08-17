"""
Reads STAR prep content from star_data.py and generates a formatted .docx file.
Run with: python3 src/best_foot_forward/utils/generate_star_prep.py
"""

import sys as _sys, os as _os
_sys.path.insert(0, (_os.environ.get('BFF_DATA_DIR') or _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '../../../data'))) + '/session')
_sys.path.insert(0, _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '..')))
del _sys, _os

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from docx_helpers import font, ACCENT_COLOR, add_border
from star_data import COMPANY, ROLE, OUTPUT_DIR, INTERVIEWER, DATE, OPS_KEY, STARS, OPS_CROSSREF
from utils.audit_log import log_event

OPS_MAP = {abbrev: name for abbrev, name in OPS_KEY}


def expand_ops(ops_str):
    """Replace OP abbreviations with full names: 'TB (desc)' → 'Think Big (desc)'"""
    parts = ops_str.split(' | ')
    expanded = []
    for part in parts:
        words = part.split(' ', 1)
        abbrev = words[0]
        if abbrev in OPS_MAP:
            rest = (' ' + words[1]) if len(words) > 1 else ''
            part = OPS_MAP[abbrev] + rest
        expanded.append(part)
    return ' | '.join(expanded)

NOTE_COLOR = RGBColor(0x88, 0x66, 0x00)  # amber for notes/callouts


def sp(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def section_header(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    font(r, bold=True, color=ACCENT_COLOR)
    sp(p, before=10, after=4)
    add_border(p)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def star_block(doc, label, items):
    """One S/T/A/R component. Label on first line, indented continuation."""
    if not items:
        return
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        if i == 0:
            # Label flush left relative to indent, then item
            p.paragraph_format.first_line_indent = Pt(-20)
            r1 = p.add_run(f"{label}  ")
            font(r1, bold=True, size=11)
            r2 = p.add_run(item)
            font(r2, size=11)
        else:
            r = p.add_run(item)
            font(r, size=11)
        sp(p, before=0, after=1)


doc = Document()
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

# ── Document header ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run(f"{COMPANY.upper()} — STAR Interview Prep")
font(r, size=14, bold=True)
sp(p, after=2)

p = doc.add_paragraph()
r = p.add_run(ROLE)
font(r, size=11, italic=True)
sp(p, after=2)

p = doc.add_paragraph()
r = p.add_run(f"Interviewer: {INTERVIEWER}  |  {DATE}")
font(r, size=10)
sp(p, after=6)

# ── OPs key ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("OPERATING PRINCIPLES KEY")
font(r, size=9, bold=True, color=ACCENT_COLOR)
sp(p, before=4, after=2)
add_border(p)

# Two-column key layout using a table
tbl = doc.add_table(rows=0, cols=4)
tbl.style = 'Table Grid'
# Suppress borders
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def remove_table_borders(table):
    tbl_el = table._tbl
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

remove_table_borders(tbl)

# Set column widths: abbrev | full name | abbrev | full name
col_widths = [Inches(0.45), Inches(2.5), Inches(0.45), Inches(2.5)]
for i, w in enumerate(col_widths):
    for cell in tbl.column_cells(i):
        cell.width = w

for i in range(0, len(OPS_KEY), 2):
    row = tbl.add_row()
    for col_offset, idx in enumerate([i, i+1] if i+1 < len(OPS_KEY) else [i]):
        abbrev, name = OPS_KEY[idx]
        cell_a = row.cells[col_offset * 2]
        cell_n = row.cells[col_offset * 2 + 1]
        p_a = cell_a.paragraphs[0]
        r_a = p_a.add_run(abbrev)
        font(r_a, size=9, bold=True, color=ACCENT_COLOR)
        p_a.paragraph_format.space_after = Pt(1)
        p_n = cell_n.paragraphs[0]
        r_n = p_n.add_run(name)
        font(r_n, size=9)
        p_n.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
sp(p, before=0, after=4)

# ── OPs cross-reference ───────────────────────────────────────────────────────
section_header(doc, "OPERATING PRINCIPLES CROSS-REFERENCE")

tbl2 = doc.add_table(rows=0, cols=2)
tbl2.style = 'Table Grid'
remove_table_borders(tbl2)

for op_label, stories in OPS_CROSSREF:
    row = tbl2.add_row()
    cell_l = row.cells[0]
    cell_r = row.cells[1]
    cell_l.width = Inches(3.0)
    cell_r.width = Inches(3.0)
    p_l = cell_l.paragraphs[0]
    r_l = p_l.add_run(op_label)
    font(r_l, size=10, bold=True, color=ACCENT_COLOR)
    p_l.paragraph_format.space_after = Pt(2)
    p_r = cell_r.paragraphs[0]
    r_r = p_r.add_run(stories)
    font(r_r, size=10)
    p_r.paragraph_format.space_after = Pt(2)

# ── STAR stories ─────────────────────────────────────────────────────────────
for story in STARS:
    page_break(doc)
    section_header(doc, story["title"])

    # Q triggers
    p = doc.add_paragraph()
    r = p.add_run("Q: " + " | ".join(story["triggers"]))
    font(r, size=10, italic=True)
    sp(p, before=2, after=4)

    # Optional note (e.g. "Direct answer — not a STAR story")
    if story.get("note"):
        p = doc.add_paragraph()
        r = p.add_run(f"NOTE: {story['note']}")
        font(r, size=10, italic=True, color=NOTE_COLOR)
        sp(p, before=0, after=3)

    # STAR components
    star_block(doc, "S", story.get("S", []))
    star_block(doc, "T", story.get("T", []))
    star_block(doc, "A", story.get("A", []))
    star_block(doc, "R", story.get("R", []))

    # Second note (pharmacy angle etc.)
    if story.get("note") and story["S"]:  # already shown above
        pass

    # Extra note field for mid-story callouts (Story 6 pharmacy angle)
    for key in story:
        if key not in ("title", "triggers", "note", "S", "T", "A", "R", "ops") and key.islower():
            p = doc.add_paragraph()
            r = p.add_run(f"► {story[key]}")
            font(r, size=10, italic=True, color=NOTE_COLOR)
            sp(p, before=4, after=2)

    # OPs line
    p = doc.add_paragraph()
    r = p.add_run(f"OPs:  {expand_ops(story['ops'])}")
    font(r, size=9, italic=True, color=ACCENT_COLOR)
    sp(p, before=4, after=8)

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(OUTPUT_DIR, exist_ok=True)
filename = f"{COMPANY}STARPrep.docx"
filepath = os.path.join(OUTPUT_DIR, filename)
doc.save(filepath)
print(f"Saved to: {filepath}")
log_event("star-prep", "generate", company=COMPANY, role=ROLE, filepath=filepath, story_count=len(STARS))
