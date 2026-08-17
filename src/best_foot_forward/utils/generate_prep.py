"""
Reads screening prep content from prep_data.py and generates a formatted .docx file.
Run with: python3 src/best_foot_forward/utils/generate_prep.py
"""

import sys as _sys, os as _os
_sys.path.insert(0, (_os.environ.get('BFF_DATA_DIR') or _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '../../../data'))) + '/session')  # data/session/ for session files
_sys.path.insert(0, _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '..')))        # src/best_foot_forward/
del _sys, _os

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_helpers import font, ACCENT_COLOR, add_border

from prep_data import COMPANY, ROLE, OUTPUT_DIR, RECRUITER, SCREEN_DATE, PREP
from utils.audit_log import log_event

# ── Formatting helpers ────────────────────────────────────────────────────────

def sp(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def section_header(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    font(r, bold=True, color=ACCENT_COLOR)
    sp(p, before=10, after=4)
    add_border(p)

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# Header
p = doc.add_paragraph()
r = p.add_run(f"{COMPANY.upper()} — Screen Prep")
font(r, size=14, bold=True)
sp(p, after=2)

p = doc.add_paragraph()
r = p.add_run(ROLE)
font(r, italic=True)
sp(p, after=2)

meta_parts = []
if RECRUITER:
    meta_parts.append(f"Interviewer: {RECRUITER}")
meta_parts.append(SCREEN_DATE)
p = doc.add_paragraph()
r = p.add_run('  |  '.join(meta_parts))
font(r)
sp(p, after=6)

# Section 1 — Tell Me About Yourself
section_header(doc, "TELL ME ABOUT YOURSELF")
p = doc.add_paragraph()
r = p.add_run(PREP['tell_me_about_yourself'])
font(r)
sp(p, after=4)

# Section 2 — Likely Questions
section_header(doc, "LIKELY QUESTIONS + BRIEF RESPONSES")
for i, qa in enumerate(PREP['questions'], 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {qa['question']}")
    font(r, bold=True)
    sp(p, before=8, after=2)

    p = doc.add_paragraph()
    r = p.add_run(qa['response'])
    font(r)
    sp(p, after=2)

# Section 3 — Questions to Ask
first_name = RECRUITER.split()[0].upper() if RECRUITER else None
ask_header = f"QUESTIONS TO ASK {first_name}" if first_name else "QUESTIONS TO ASK"
section_header(doc, ask_header)
for i, q in enumerate(PREP['questions_to_ask'], 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {q}")
    font(r)
    sp(p, before=6, after=4)

# ── Save ──────────────────────────────────────────────────────────────────────

os.makedirs(OUTPUT_DIR, exist_ok=True)
filename = f"{COMPANY}ScreenPrep.docx"
filepath = os.path.join(OUTPUT_DIR, filename)
doc.save(filepath)
print(f"Saved to: {filepath}")
log_event("screening-prep", "generate", company=COMPANY, role=ROLE, filepath=filepath)
