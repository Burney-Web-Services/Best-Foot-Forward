"""
Reads cover letter content from letter_data.py and generates a formatted .docx file.
Contact info is read from SQLite. Output path is the directory containing the
source JD file (JD_FILE_PATH in letter_data.py).
Run with: python3 src/best_foot_forward/utils/generate_letter.py
"""

import sys as _sys, os as _os
_sys.path.insert(0, (_os.environ.get('BFF_DATA_DIR') or _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '../../../data'))) + '/session')  # data/session/ for session files
_sys.path.insert(0, _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '..')))        # src/best_foot_forward/
del _sys, _os

import os
from datetime import date, datetime, timezone
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_helpers import font, get_accent_color, compose_signoff

from db import get_contact, get_conn, register_file, resolve_jd_path, resolve_or_create_jd
from utils.audit_log import log_event
from letter_data import COMPANY, ROLE, JD_FILE_PATH, LETTER
JD_FILE_PATH = resolve_jd_path(JD_FILE_PATH)
try:
    from resume_data import LOCATION_OVERRIDE
except ImportError:
    LOCATION_OVERRIDE = None

# ── Formatting helpers ────────────────────────────────────────────────────────

def sp(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8,
             bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    font(r, size=size, bold=bold, italic=italic)
    sp(p, before=before, after=after)
    return p

def render_letter_as_text(letter, contact_info, location_override=None):
    lines = []
    lines.append(contact_info['name'].upper())
    contact_parts = [contact_info.get('phone', ''), contact_info.get('email', ''), location_override or contact_info.get('location', '')]
    contact_line = '  |  '.join(part for part in contact_parts if part)
    lines.append(contact_line)
    lines.append('')
    lines.append(date.today().strftime('%B %-d, %Y'))
    lines.append('')
    lines.append(letter['salutation'])
    lines.append('')
    for para in letter['paragraphs']:
        lines.append(para)
        lines.append('')
    lines.extend(compose_signoff(letter.get('closing'), contact_info['name']))
    return '\n'.join(lines)

# ── Data ──────────────────────────────────────────────────────────────────────

CONTACT_INFO = get_contact()

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(CONTACT_INFO['name'].upper())
font(r, size=16, bold=True, color=get_accent_color())
sp(p, after=2)

# Contact line
contact_parts = [CONTACT_INFO.get('phone', ''), CONTACT_INFO.get('email', ''), LOCATION_OVERRIDE or CONTACT_INFO.get('location', '')]
contact_line = '  |  '.join(part for part in contact_parts if part)
add_para(doc, contact_line, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

# Date
formatted_date = date.today().strftime('%B %-d, %Y')
add_para(doc, formatted_date, after=12)

# Salutation
add_para(doc, LETTER['salutation'], after=8)

# Body paragraphs
for para_text in LETTER['paragraphs']:
    add_para(doc, para_text, after=8)

# Closing / Signature — the name comes from CONTACT_INFO unless the closing already
# carries it (conventions have differed); see compose_signoff.
signoff = compose_signoff(LETTER.get('closing'), CONTACT_INFO['name'])
for i, block in enumerate(signoff):
    last = i == len(signoff) - 1
    add_para(doc, block, before=4 if i == 0 else 0, after=0 if last else 24)

# ── Save ──────────────────────────────────────────────────────────────────────

output_dir = os.path.dirname(JD_FILE_PATH)
os.makedirs(output_dir, exist_ok=True)
# sanitize path separators in the role/company (e.g. "Data/Integration")
filename = f"{CONTACT_INFO['name']} Cover Letter - {COMPANY} {ROLE}.docx".replace("/", "-").replace("\\", "-")
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved to: {filepath}")

# Generate plain-text version with same filename stem
txt_filename = filename.replace(".docx", ".txt")
txt_filepath = os.path.join(output_dir, txt_filename)
txt_content = render_letter_as_text(LETTER, CONTACT_INFO, LOCATION_OVERRIDE)
with open(txt_filepath, 'w', encoding='utf-8') as f:
    f.write(txt_content)
print(f"Saved to: {txt_filepath}")

# ── Persist letter to SQLite ──────────────────────────────────────────────────

conn = get_conn()
# Same resolve-or-create as generate_resume.py: a letter written for a job that was
# never evaluated would otherwise insert an application with jd_id NULL.
jd_id, _jd_action = resolve_or_create_jd(conn, COMPANY, ROLE, JD_FILE_PATH)
if _jd_action == "created":
    print(f"[generate_letter] No jds row for {COMPANY} - {ROLE}; inserted one (jd_id={jd_id}).")

# jd_id is never None now, so this is a plain equality match. It previously used
# `IS ?`, which with a NULL jd_id matched *any* orphaned application and would
# attach this letter to an unrelated company's row.
app = conn.execute(
    "SELECT id FROM applications WHERE jd_id = ?",
    (jd_id,)
).fetchone()

if app:
    app_id = app["id"]
    conn.execute(
        "UPDATE applications SET letter_salutation=?, letter_closing=? WHERE id=?",
        (LETTER['salutation'], LETTER['closing'], app_id)
    )
    conn.execute("DELETE FROM application_letter_paragraphs WHERE application_id=?", (app_id,))
    for i, body in enumerate(LETTER['paragraphs']):
        conn.execute(
            "INSERT INTO application_letter_paragraphs (application_id, position, body) VALUES (?,?,?)",
            (app_id, i, body)
        )
else:
    # Writing letter files means the seeker applied, so stamp status/applied_at
    # here rather than leaning on the schema default for status (which would
    # mark the row applied while leaving applied_at NULL). Without this, when the
    # letter generator creates the row before generate_resume/track_application,
    # both of those find an existing row and neither back-fills applied_at, so the
    # application stays invisible to every report that filters on applied_at.
    cur = conn.execute(
        "INSERT INTO applications (jd_id, created_at, status, applied_at, resume_summary, letter_salutation, letter_closing) "
        "VALUES (?, ?, 'applied', ?, '', ?, ?)",
        (jd_id, datetime.now(timezone.utc).isoformat(), date.today().isoformat(), LETTER['salutation'], LETTER['closing'])
    )
    app_id = cur.lastrowid
    for i, body in enumerate(LETTER['paragraphs']):
        conn.execute(
            "INSERT INTO application_letter_paragraphs (application_id, position, body) VALUES (?,?,?)",
            (app_id, i, body)
        )

conn.commit()
conn.close()

log_event("resume-tailor", "generate_letter", application_id=app_id, jd_id=jd_id, company=COMPANY, role=ROLE, filepath=filepath)

register_file(
    filepath, "letter",
    f"Cover letter for {COMPANY} – {ROLE}",
    jd_id=jd_id, application_id=app_id,
)

register_file(
    txt_filepath, "letter",
    f"Cover letter for {COMPANY} – {ROLE} (plain text)",
    jd_id=jd_id, application_id=app_id,
)
